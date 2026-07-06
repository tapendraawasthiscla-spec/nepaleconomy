"""
DOCX handler utilizing python-docx.
"""
import io
from typing import Dict, Any
import docx

from app.legacy_fonts.converter import convert_legacy_text, is_legacy_font

def extract_docx(docx_bytes: bytes) -> Dict[str, Any]:
    """
    Extracts text from DOCX files and converts legacy fonts dynamically.
    """
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
    except Exception as e:
        raise ValueError(f"Failed to read DOCX file: {e}")

    detected_fonts = set()
    had_legacy = False
    text_blocks = []

    def process_paragraph(p) -> str:
        nonlocal had_legacy
        
        para_font = None
        if p.style and p.style.font and p.style.font.name:
            para_font = p.style.font.name
            
        para_text = []
        for run in p.runs:
            raw_text = run.text
            if not raw_text:
                continue
                
            font_name = run.font.name if run.font and run.font.name else para_font
            
            if font_name is None:
                detected_fonts.add("font-unknown")
                para_text.append(raw_text)
            else:
                detected_fonts.add(font_name)
                if is_legacy_font(font_name):
                    had_legacy = True
                    para_text.append(convert_legacy_text(raw_text, font_name))
                else:
                    para_text.append(raw_text)
        return "".join(para_text)

    for p in doc.paragraphs:
        text_blocks.append(process_paragraph(p))
        
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = []
                for p in cell.paragraphs:
                    cell_text.append(process_paragraph(p))
                row_data.append(" ".join(cell_text).strip())
            text_blocks.append(" | ".join(row_data))
            
    final_text = "\n".join(text_blocks).strip()
    
    return {
        "text": final_text,
        "had_legacy_fonts": had_legacy,
        "detected_fonts": list(detected_fonts)
    }
